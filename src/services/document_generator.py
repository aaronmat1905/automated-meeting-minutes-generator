"""
Document Generator Service
Generates meeting minutes in multiple formats (PDF, Markdown, Text, DOCX)
Supports multiple templates: MRS, MTQP, MSAD
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import json

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from src.config import Config

logger = logging.getLogger(__name__)


class DocumentGenerationError(Exception):
    """Custom exception for document generation errors"""
    pass


class DocumentGenerator:
    """Service for generating meeting minutes documents in various formats"""

    def __init__(self, config: Config):
        self.config = config
        self.export_folder = config.EXPORT_FOLDER
        self.export_folder.mkdir(parents=True, exist_ok=True)

    def generate_minutes(
        self,
        meeting_data: Dict,
        transcript_analysis: Dict,
        template: str = 'MRS',
        formats: Optional[List[str]] = None
    ) -> Dict[str, str]:
        """
        Generate meeting minutes in specified formats

        Args:
            meeting_data: Meeting metadata and information
            transcript_analysis: Analysis results from Gemini agent
            template: Template type ('MRS', 'MTQP', 'MSAD')
            formats: List of formats to generate (default: all)

        Returns:
            Dictionary mapping format to file path
        """
        if formats is None:
            formats = self.config.EXPORT_FORMATS

        template = template.upper()
        if template not in ['MRS', 'MTQP', 'MSAD']:
            raise DocumentGenerationError(f"Invalid template: {template}")

        # Generate content based on template
        if template == 'MRS':
            content = self._generate_mrs_content(meeting_data, transcript_analysis)
        elif template == 'MTQP':
            content = self._generate_mtqp_content(meeting_data, transcript_analysis)
        else:  # MSAD
            content = self._generate_msad_content(meeting_data, transcript_analysis)

        # Generate files in requested formats
        output_files = {}

        try:
            base_filename = self._generate_filename(meeting_data, template)

            if 'pdf' in formats:
                pdf_path = self._generate_pdf(content, base_filename)
                output_files['pdf'] = pdf_path

            if 'markdown' in formats:
                md_path = self._generate_markdown(content, base_filename)
                output_files['markdown'] = md_path

            if 'txt' in formats:
                txt_path = self._generate_text(content, base_filename)
                output_files['txt'] = txt_path

            if 'docx' in formats:
                docx_path = self._generate_docx(content, base_filename)
                output_files['docx'] = docx_path

            logger.info(f"Generated minutes in {len(output_files)} formats")
            return output_files

        except Exception as e:
            logger.error(f"Error generating minutes: {str(e)}")
            raise DocumentGenerationError(f"Failed to generate minutes: {str(e)}")

    def _generate_mrs_content(self, meeting_data: Dict, analysis: Dict) -> Dict:
        """
        Generate MRS (Meeting Recording System) format content

        MRS Format:
        - Meeting Information
        - Attendees
        - Agenda Items
        - Discussion Summary
        - Action Items
        - Decisions Made
        """
        content = {
            'template': 'MRS',
            'title': 'Meeting Recording System (MRS) Minutes',
            'sections': []
        }

        # Meeting Information
        content['sections'].append({
            'heading': 'Meeting Information',
            'type': 'info',
            'items': [
                ('Meeting Title', meeting_data.get('title', 'N/A')),
                ('Date', meeting_data.get('date', datetime.now().strftime('%Y-%m-%d'))),
                ('Time', meeting_data.get('time', 'N/A')),
                ('Location/Platform', meeting_data.get('location', 'Virtual')),
                ('Organizer', meeting_data.get('organizer', 'N/A')),
                ('Duration', meeting_data.get('duration', 'N/A')),
            ]
        })

        # Attendees
        content['sections'].append({
            'heading': 'Attendees',
            'type': 'list',
            'items': meeting_data.get('participants', [])
        })

        # Agenda
        if meeting_data.get('agenda'):
            content['sections'].append({
                'heading': 'Agenda',
                'type': 'text',
                'content': meeting_data['agenda']
            })

        # Discussion Summary
        if analysis.get('key_topics'):
            content['sections'].append({
                'heading': 'Discussion Summary',
                'type': 'topics',
                'items': analysis['key_topics']
            })

        # Action Items
        if analysis.get('action_items'):
            content['sections'].append({
                'heading': 'Action Items',
                'type': 'action_items',
                'items': analysis['action_items']
            })

        # Decisions Made
        if analysis.get('decisions'):
            content['sections'].append({
                'heading': 'Decisions Made',
                'type': 'decisions',
                'items': analysis['decisions']
            })

        # Open Questions
        if analysis.get('open_questions'):
            content['sections'].append({
                'heading': 'Open Questions / Follow-up Items',
                'type': 'questions',
                'items': analysis['open_questions']
            })

        return content

    def _generate_mtqp_content(self, meeting_data: Dict, analysis: Dict) -> Dict:
        """
        Generate MTQP (Meeting Topic, Questions, and Points) format content

        MTQP Format:
        - Executive Summary
        - Meeting Topics
        - Key Discussion Points
        - Questions Raised
        - Action Points
        - Next Steps
        """
        content = {
            'template': 'MTQP',
            'title': 'Meeting Topics, Questions, and Points (MTQP)',
            'sections': []
        }

        # Meeting Header
        content['sections'].append({
            'heading': 'Meeting Information',
            'type': 'info',
            'items': [
                ('Title', meeting_data.get('title', 'N/A')),
                ('Date', meeting_data.get('date', datetime.now().strftime('%Y-%m-%d'))),
                ('Participants', ', '.join(meeting_data.get('participants', []))),
            ]
        })

        # Executive Summary
        if analysis.get('executive_summary'):
            summary = analysis['executive_summary']
            content['sections'].append({
                'heading': 'Executive Summary',
                'type': 'text',
                'content': summary.get('overview', 'N/A')
            })

        # Meeting Topics
        if analysis.get('key_topics'):
            content['sections'].append({
                'heading': 'Topics Discussed',
                'type': 'topics',
                'items': analysis['key_topics']
            })

        # Questions Raised
        if analysis.get('open_questions'):
            content['sections'].append({
                'heading': 'Questions Raised',
                'type': 'questions',
                'items': analysis['open_questions']
            })

        # Key Points and Decisions
        if analysis.get('decisions'):
            content['sections'].append({
                'heading': 'Key Points and Decisions',
                'type': 'decisions',
                'items': analysis['decisions']
            })

        # Action Points
        if analysis.get('action_items'):
            content['sections'].append({
                'heading': 'Action Points',
                'type': 'action_items',
                'items': analysis['action_items']
            })

        return content

    def _generate_msad_content(self, meeting_data: Dict, analysis: Dict) -> Dict:
        """
        Generate MSAD (Meeting Summary and Action Dashboard) format content

        MSAD Format:
        - Quick Overview
        - Key Outcomes
        - Action Dashboard (with status tracking)
        - Decisions Log
        - Risk and Blockers
        """
        content = {
            'template': 'MSAD',
            'title': 'Meeting Summary and Action Dashboard (MSAD)',
            'sections': []
        }

        # Quick Overview
        content['sections'].append({
            'heading': 'Quick Overview',
            'type': 'info',
            'items': [
                ('Meeting', meeting_data.get('title', 'N/A')),
                ('Date', meeting_data.get('date', datetime.now().strftime('%Y-%m-%d'))),
                ('Duration', meeting_data.get('duration', 'N/A')),
                ('Attendees', ', '.join(meeting_data.get('participants', [])[:5]) + '...' if len(meeting_data.get('participants', [])) > 5 else ', '.join(meeting_data.get('participants', []))),
            ]
        })

        # Executive Summary
        if analysis.get('executive_summary'):
            summary = analysis['executive_summary']

            if summary.get('key_outcomes'):
                content['sections'].append({
                    'heading': 'Key Outcomes',
                    'type': 'list',
                    'items': summary['key_outcomes']
                })

        # Action Dashboard
        if analysis.get('action_items'):
            # Add status field to action items for tracking
            action_items_with_status = []
            for item in analysis['action_items']:
                item_copy = item.copy()
                item_copy['status'] = 'Pending'  # Default status
                action_items_with_status.append(item_copy)

            content['sections'].append({
                'heading': 'Action Dashboard',
                'type': 'action_dashboard',
                'items': action_items_with_status
            })

        # Decisions Log
        if analysis.get('decisions'):
            content['sections'].append({
                'heading': 'Decisions Log',
                'type': 'decisions',
                'items': analysis['decisions']
            })

        # Risks and Blockers
        if analysis.get('executive_summary', {}).get('risks_or_blockers'):
            risks = analysis['executive_summary']['risks_or_blockers']
            if risks:
                content['sections'].append({
                    'heading': 'Risks and Blockers',
                    'type': 'list',
                    'items': risks if isinstance(risks, list) else [risks]
                })

        # Next Meeting
        if analysis.get('executive_summary', {}).get('next_meeting'):
            content['sections'].append({
                'heading': 'Next Steps',
                'type': 'text',
                'content': analysis['executive_summary']['next_meeting']
            })

        return content

    def _generate_pdf(self, content: Dict, base_filename: str) -> str:
        """Generate PDF document"""
        filename = f"{base_filename}.pdf"
        filepath = self.export_folder / filename

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(content['title'], title_style))
        story.append(Spacer(1, 12))

        # Add sections
        for section in content['sections']:
            # Section heading
            story.append(Paragraph(section['heading'], styles['Heading2']))
            story.append(Spacer(1, 12))

            # Section content
            if section['type'] == 'info':
                data = [[k, v] for k, v in section['items']]
                table = Table(data, colWidths=[2*inch, 4*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey)
                ]))
                story.append(table)

            elif section['type'] == 'list':
                for item in section['items']:
                    story.append(Paragraph(f"• {item}", styles['Normal']))

            elif section['type'] == 'text':
                story.append(Paragraph(section['content'], styles['Normal']))

            elif section['type'] == 'action_items' or section['type'] == 'action_dashboard':
                for item in section['items']:
                    action_text = f"<b>{item.get('description', 'N/A')}</b><br/>"
                    action_text += f"Owner: {item.get('owner', 'Unassigned')}<br/>"
                    action_text += f"Due: {item.get('due_date', 'Not specified')}<br/>"
                    action_text += f"Priority: {item.get('priority', 'Medium').title()}"
                    if section['type'] == 'action_dashboard':
                        action_text += f"<br/>Status: {item.get('status', 'Pending')}"
                    story.append(Paragraph(action_text, styles['Normal']))
                    story.append(Spacer(1, 12))

            elif section['type'] == 'decisions':
                for item in section['items']:
                    decision_text = f"<b>{item.get('decision', 'N/A')}</b><br/>"
                    if item.get('rationale'):
                        decision_text += f"Rationale: {item['rationale']}<br/>"
                    if item.get('stakeholders'):
                        stakeholders = ', '.join(item['stakeholders']) if isinstance(item['stakeholders'], list) else item['stakeholders']
                        decision_text += f"Stakeholders: {stakeholders}"
                    story.append(Paragraph(decision_text, styles['Normal']))
                    story.append(Spacer(1, 12))

            elif section['type'] == 'topics':
                for item in section['items']:
                    topic_text = f"<b>{item.get('topic', 'N/A')}</b><br/>"
                    topic_text += f"{item.get('summary', 'N/A')}"
                    story.append(Paragraph(topic_text, styles['Normal']))
                    story.append(Spacer(1, 12))

            elif section['type'] == 'questions':
                for item in section['items']:
                    question_text = f"<b>Q:</b> {item.get('question', 'N/A')}<br/>"
                    if item.get('who_needs_to_answer'):
                        question_text += f"Owner: {item['who_needs_to_answer']}"
                    story.append(Paragraph(question_text, styles['Normal']))
                    story.append(Spacer(1, 12))

            story.append(Spacer(1, 20))

        # Build PDF
        doc.build(story)
        logger.info(f"PDF generated: {filepath}")
        return str(filepath)

    def _generate_markdown(self, content: Dict, base_filename: str) -> str:
        """Generate Markdown document"""
        filename = f"{base_filename}.md"
        filepath = self.export_folder / filename

        md_content = [f"# {content['title']}\n"]

        for section in content['sections']:
            md_content.append(f"\n## {section['heading']}\n")

            if section['type'] == 'info':
                for k, v in section['items']:
                    md_content.append(f"**{k}:** {v}  ")

            elif section['type'] == 'list':
                for item in section['items']:
                    md_content.append(f"- {item}")

            elif section['type'] == 'text':
                md_content.append(f"\n{section['content']}\n")

            elif section['type'] in ['action_items', 'action_dashboard']:
                md_content.append("\n| Description | Owner | Due Date | Priority | Status |")
                md_content.append("|-------------|-------|----------|----------|--------|")
                for item in section['items']:
                    desc = item.get('description', 'N/A')
                    owner = item.get('owner', 'Unassigned')
                    due = item.get('due_date', 'Not specified')
                    priority = item.get('priority', 'Medium').title()
                    status = item.get('status', 'N/A') if section['type'] == 'action_dashboard' else '-'
                    md_content.append(f"| {desc} | {owner} | {due} | {priority} | {status} |")

            elif section['type'] == 'decisions':
                for i, item in enumerate(section['items'], 1):
                    md_content.append(f"\n### Decision {i}: {item.get('decision', 'N/A')}")
                    if item.get('rationale'):
                        md_content.append(f"**Rationale:** {item['rationale']}")
                    if item.get('stakeholders'):
                        stakeholders = ', '.join(item['stakeholders']) if isinstance(item['stakeholders'], list) else item['stakeholders']
                        md_content.append(f"**Stakeholders:** {stakeholders}")

            elif section['type'] == 'topics':
                for i, item in enumerate(section['items'], 1):
                    md_content.append(f"\n### {i}. {item.get('topic', 'N/A')}")
                    md_content.append(f"{item.get('summary', 'N/A')}")

            elif section['type'] == 'questions':
                for i, item in enumerate(section['items'], 1):
                    md_content.append(f"\n**Q{i}:** {item.get('question', 'N/A')}")
                    if item.get('who_needs_to_answer'):
                        md_content.append(f"*Owner: {item['who_needs_to_answer']}*")

            md_content.append("\n")

        # Add footer
        md_content.append(f"\n---\n*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_content))

        logger.info(f"Markdown generated: {filepath}")
        return str(filepath)

    def _generate_text(self, content: Dict, base_filename: str) -> str:
        """Generate plain text document"""
        filename = f"{base_filename}.txt"
        filepath = self.export_folder / filename

        text_content = [f"{content['title']}\n{'=' * len(content['title'])}\n"]

        for section in content['sections']:
            text_content.append(f"\n{section['heading']}\n{'-' * len(section['heading'])}")

            if section['type'] == 'info':
                for k, v in section['items']:
                    text_content.append(f"{k}: {v}")

            elif section['type'] == 'list':
                for item in section['items']:
                    text_content.append(f"  • {item}")

            elif section['type'] == 'text':
                text_content.append(f"\n{section['content']}\n")

            elif section['type'] in ['action_items', 'action_dashboard']:
                for i, item in enumerate(section['items'], 1):
                    text_content.append(f"\n{i}. {item.get('description', 'N/A')}")
                    text_content.append(f"   Owner: {item.get('owner', 'Unassigned')}")
                    text_content.append(f"   Due Date: {item.get('due_date', 'Not specified')}")
                    text_content.append(f"   Priority: {item.get('priority', 'Medium').title()}")
                    if section['type'] == 'action_dashboard':
                        text_content.append(f"   Status: {item.get('status', 'Pending')}")

            elif section['type'] == 'decisions':
                for i, item in enumerate(section['items'], 1):
                    text_content.append(f"\n{i}. {item.get('decision', 'N/A')}")
                    if item.get('rationale'):
                        text_content.append(f"   Rationale: {item['rationale']}")

            elif section['type'] == 'topics':
                for i, item in enumerate(section['items'], 1):
                    text_content.append(f"\n{i}. {item.get('topic', 'N/A')}")
                    text_content.append(f"   {item.get('summary', 'N/A')}")

            elif section['type'] == 'questions':
                for i, item in enumerate(section['items'], 1):
                    text_content.append(f"\n{i}. {item.get('question', 'N/A')}")
                    if item.get('who_needs_to_answer'):
                        text_content.append(f"   Owner: {item['who_needs_to_answer']}")

            text_content.append("")

        # Add footer
        text_content.append(f"\nGenerated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))

        logger.info(f"Text file generated: {filepath}")
        return str(filepath)

    def _generate_docx(self, content: Dict, base_filename: str) -> str:
        """Generate DOCX document"""
        filename = f"{base_filename}.docx"
        filepath = self.export_folder / filename

        doc = Document()

        # Title
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        for section in content['sections']:
            doc.add_heading(section['heading'], level=1)

            if section['type'] == 'info':
                table = doc.add_table(rows=len(section['items']), cols=2)
                table.style = 'Light Grid Accent 1'
                for i, (k, v) in enumerate(section['items']):
                    table.rows[i].cells[0].text = str(k)
                    table.rows[i].cells[1].text = str(v)

            elif section['type'] == 'list':
                for item in section['items']:
                    doc.add_paragraph(str(item), style='List Bullet')

            elif section['type'] == 'text':
                doc.add_paragraph(section['content'])

            elif section['type'] in ['action_items', 'action_dashboard']:
                for item in section['items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item.get('description', 'N/A')).bold = True
                    p.add_run(f"\nOwner: {item.get('owner', 'Unassigned')}")
                    p.add_run(f"\nDue: {item.get('due_date', 'Not specified')}")
                    p.add_run(f"\nPriority: {item.get('priority', 'Medium').title()}")
                    if section['type'] == 'action_dashboard':
                        p.add_run(f"\nStatus: {item.get('status', 'Pending')}")

            elif section['type'] == 'decisions':
                for item in section['items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item.get('decision', 'N/A')).bold = True
                    if item.get('rationale'):
                        p.add_run(f"\nRationale: {item['rationale']}")

            elif section['type'] == 'topics':
                for item in section['items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item.get('topic', 'N/A')).bold = True
                    p.add_run(f"\n{item.get('summary', 'N/A')}")

            elif section['type'] == 'questions':
                for item in section['items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"Q: {item.get('question', 'N/A')}")
                    if item.get('who_needs_to_answer'):
                        p.add_run(f"\nOwner: {item['who_needs_to_answer']}")

        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(str(filepath))
        logger.info(f"DOCX generated: {filepath}")
        return str(filepath)

    def _generate_filename(self, meeting_data: Dict, template: str) -> str:
        """Generate filename for exported document"""
        title = meeting_data.get('title', 'meeting')
        # Clean title for filename
        title = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in title)
        title = title.replace(' ', '_')[:50]  # Limit length

        date = meeting_data.get('date', datetime.now().strftime('%Y%m%d'))
        timestamp = datetime.now().strftime('%H%M%S')

        return f"{title}_{template}_{date}_{timestamp}"
